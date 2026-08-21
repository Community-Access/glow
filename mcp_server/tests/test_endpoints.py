"""
Unit tests for GLOW MCP Server endpoints using FastAPI TestClient.
"""
from pathlib import Path
from fastapi.testclient import TestClient
import pytest



import sys
# Add project root and mcp_server to sys.path
project_root = Path(__file__).parent.parent.parent.resolve()
sys.path.insert(0, str(project_root))
sys.path.insert(0, str(project_root / "mcp_server"))
from main import app  # noqa: E402 - import requires the sys.path setup above
import main as mcp_main  # noqa: E402 - import requires the sys.path setup above

client = TestClient(app)


def test_health():
    resp = client.get("/health")
    assert resp.status_code == 200
    assert resp.json()["status"] == "ok"


def test_audit_markdown(tmp_path):
    md_file = tmp_path / "test.md"
    md_file.write_text("# Heading\n\nSome text.")
    with md_file.open("rb") as f:
        resp = client.post("/audit", files={"file": ("test.md", f, "text/markdown")}, data={"format": "markdown"})
    assert resp.status_code == 200
    assert "score" in resp.text or "result" in resp.text


def test_audit_docx(tmp_path):
    # Minimal DOCX file (could use python-docx to generate, or skip if not available)
    try:
        from docx import Document
    except ImportError:
        pytest.skip("python-docx not installed")
    docx_file = tmp_path / "test.docx"
    doc = Document()
    doc.add_heading("Test Heading", 0)
    doc.add_paragraph("Some text.")
    doc.save(docx_file)
    with docx_file.open("rb") as f:
        resp = client.post("/audit", files={"file": ("test.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}, data={"format": "docx"})
    assert resp.status_code == 200
    assert "score" in resp.text or "result" in resp.text


def test_page_flow_extract_success(monkeypatch):
    monkeypatch.setattr(
        mcp_main,
        "run_page_flow_extract",
        lambda source_url, max_pages=5, follow_pagination=True: {
            "source_url": source_url,
            "final_url": source_url,
            "title": "Fixture Title",
            "text": "Fixture extracted body.",
            "page_urls": [source_url],
            "page_count": 1,
            "char_count": 23,
        },
    )

    resp = client.post(
        "/page-flow",
        data={
            "source_url": "https://example.com/article",
            "max_pages": "3",
            "follow_pagination": "true",
        },
    )

    assert resp.status_code == 200
    body = resp.json()
    assert body["title"] == "Fixture Title"
    assert body["page_count"] == 1
    assert "Fixture extracted body" in body["text"]


def test_page_flow_extract_error(monkeypatch):
    def _raise(*args, **kwargs):
        raise ValueError("Please enter a valid article URL.")

    monkeypatch.setattr(mcp_main, "run_page_flow_extract", _raise)

    resp = client.post(
        "/page-flow",
        data={
            "source_url": "https://example.com/bad",
        },
    )

    assert resp.status_code == 400
    assert "error" in resp.json()

def test_report_markdown_json_and_text(tmp_path):
    md_file = tmp_path / "report.md"
    md_file.write_text("# Heading\n\nSome text.")
    for report_type, marker in (("json", "score"), ("text", "GLOW accessibility audit")):
        with md_file.open("rb") as f:
            resp = client.post(
                "/report",
                files={"file": ("report.md", f, "text/markdown")},
                data={"format": "markdown", "report_type": report_type},
            )
        assert resp.status_code == 200, resp.text
        assert marker in resp.json()["report"]


def test_fix_docx_returns_json_safe_payload(tmp_path):
    try:
        from docx import Document
    except ImportError:
        pytest.skip("python-docx not installed")
    docx_file = tmp_path / "fixme.docx"
    doc = Document()
    doc.add_heading("Heading", 0)
    doc.add_paragraph("Some text.")
    doc.save(docx_file)
    with docx_file.open("rb") as f:
        resp = client.post(
            "/fix",
            files={"file": ("fixme.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            data={"format": "docx"},
        )
    assert resp.status_code == 200, resp.text
    body = resp.json()
    assert set(body) == {
        "fixed_file",
        "fixed_file_base64",
        "total_fixes",
        "fix_records",
        "post_fix_audit",
        "warnings",
    }
    assert isinstance(body["total_fixes"], int)
    assert "score" in body["post_fix_audit"]
    # No server-local paths may leak; the artifact ships inline as base64.
    assert "/" not in body["fixed_file"] and "\\" not in body["fixed_file"]
    import base64
    fixed_bytes = base64.b64decode(body["fixed_file_base64"])
    assert fixed_bytes[:2] == b"PK"  # docx is a zip container


def test_convert_endpoint_plumbing(monkeypatch, tmp_path):
    # Conversion itself needs pandoc; this validates the endpoint contract.
    import glow_mcp_utils

    def _fake_convert(file_path, from_fmt, to_fmt, output_path=None):
        out = tmp_path / "out.md"
        out.write_text("# Converted")
        return out, "# Converted"

    monkeypatch.setattr(mcp_main, "run_convert", _fake_convert, raising=False)
    monkeypatch.setattr(glow_mcp_utils, "run_convert", _fake_convert)
    src = tmp_path / "in.docx"
    src.write_bytes(b"stub")
    with src.open("rb") as f:
        resp = client.post(
            "/convert",
            files={"file": ("in.docx", f, "application/octet-stream")},
            data={"from_format": "docx", "to_format": "markdown"},
        )
    assert resp.status_code == 200, resp.text
    body = resp.json()
    assert body["converted_text"] == "# Converted"
