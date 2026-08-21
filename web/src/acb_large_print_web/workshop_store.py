"""Persistence helpers for Workshop Mode (7.3.0).

This module provides a lightweight SQLite store under instance/ so workshop
sessions, submissions, and peer feedback can be used in conference/training
environments without additional infrastructure.
"""

from __future__ import annotations

import hashlib
import json
import os
import re
import secrets
import sqlite3
from datetime import UTC, datetime, timedelta
from html import escape
from io import BytesIO
from pathlib import Path

from flask import current_app

_SESSION_CODE_RE = re.compile(r"^[a-zA-Z0-9_-]{3,64}$")


def _utc_now() -> str:
    return datetime.now(UTC).isoformat()


def _db_path() -> Path:
    p = Path(current_app.instance_path)
    p.mkdir(parents=True, exist_ok=True)
    return p / "workshop_mode.db"


def _conn() -> sqlite3.Connection:
    conn = sqlite3.connect(str(_db_path()))
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA journal_mode=WAL")
    _ensure_schema(conn)
    return conn


def _ensure_schema(conn: sqlite3.Connection) -> None:
    conn.execute(
        "CREATE TABLE IF NOT EXISTS workshop_sessions ("
        " session_code TEXT PRIMARY KEY,"
        " title TEXT NOT NULL,"
        " event_name TEXT,"
        " created_at_utc TEXT NOT NULL"
        ")"
    )
    conn.execute(
        "CREATE TABLE IF NOT EXISTS workshop_submissions ("
        " id INTEGER PRIMARY KEY AUTOINCREMENT,"
        " session_code TEXT NOT NULL,"
        " activity_key TEXT NOT NULL,"
        " participant_key TEXT,"
        " display_name TEXT NOT NULL,"
        " anonymity_mode INTEGER NOT NULL DEFAULT 0,"
        " content_text TEXT NOT NULL,"
        " content_json TEXT,"
        " is_draft INTEGER NOT NULL DEFAULT 0,"
        " created_at_utc TEXT NOT NULL,"
        " updated_at_utc TEXT NOT NULL,"
        " FOREIGN KEY(session_code) REFERENCES workshop_sessions(session_code)"
        ")"
    )
    conn.execute(
        "CREATE TABLE IF NOT EXISTS workshop_participants ("
        " participant_key TEXT PRIMARY KEY,"
        " session_code TEXT NOT NULL,"
        " display_name TEXT NOT NULL,"
        " login_email TEXT,"
        " created_at_utc TEXT NOT NULL,"
        " updated_at_utc TEXT NOT NULL,"
        " last_seen_at_utc TEXT NOT NULL,"
        " FOREIGN KEY(session_code) REFERENCES workshop_sessions(session_code)"
        ")"
    )
    conn.execute(
        "CREATE TABLE IF NOT EXISTS workshop_conference_codes ("
        " access_code TEXT PRIMARY KEY,"
        " session_code TEXT NOT NULL,"
        " session_title TEXT NOT NULL,"
        " event_name TEXT,"
        " active INTEGER NOT NULL DEFAULT 1,"
        " created_at_utc TEXT NOT NULL,"
        " updated_at_utc TEXT NOT NULL"
        ")"
    )
    conn.execute(
        "CREATE TABLE IF NOT EXISTS workshop_nudges ("
        " session_code TEXT NOT NULL,"
        " participant_key TEXT NOT NULL,"
        " sent_at_utc TEXT NOT NULL,"
        " PRIMARY KEY (session_code, participant_key)"
        ")"
    )
    conn.execute(
        "CREATE TABLE IF NOT EXISTS workshop_return_links ("
        " token_hash TEXT PRIMARY KEY,"
        " session_code TEXT NOT NULL,"
        " participant_key TEXT NOT NULL,"
        " created_at_utc TEXT NOT NULL,"
        " expires_at_utc TEXT NOT NULL,"
        " used_at_utc TEXT"
        ")"
    )
    conn.execute(
        "CREATE TABLE IF NOT EXISTS workshop_feedback ("
        " id INTEGER PRIMARY KEY AUTOINCREMENT,"
        " session_code TEXT NOT NULL,"
        " submission_id INTEGER NOT NULL,"
        " reviewer_display_name TEXT NOT NULL,"
        " strength TEXT NOT NULL,"
        " risk_or_gap TEXT NOT NULL,"
        " recommended_safeguard TEXT NOT NULL,"
        " reuse_suggestion TEXT NOT NULL,"
        " created_at_utc TEXT NOT NULL,"
        " FOREIGN KEY(session_code) REFERENCES workshop_sessions(session_code),"
        " FOREIGN KEY(submission_id) REFERENCES workshop_submissions(id)"
        ")"
    )
    conn.execute(
        "CREATE TABLE IF NOT EXISTS workshop_follow_through ("
        " id INTEGER PRIMARY KEY AUTOINCREMENT,"
        " session_code TEXT NOT NULL,"
        " source_submission_id INTEGER,"
        " kind TEXT NOT NULL,"
        " title TEXT NOT NULL,"
        " details TEXT NOT NULL,"
        " owner_name TEXT NOT NULL,"
        " due_date TEXT,"
        " status TEXT NOT NULL DEFAULT 'open',"
        " created_at_utc TEXT NOT NULL,"
        " updated_at_utc TEXT NOT NULL,"
        " FOREIGN KEY(session_code) REFERENCES workshop_sessions(session_code),"
        " FOREIGN KEY(source_submission_id) REFERENCES workshop_submissions(id)"
        ")"
    )
    # Schema evolution for existing deployments.
    cols = [str(r["name"]) for r in conn.execute("PRAGMA table_info(workshop_submissions)").fetchall()]
    if "participant_key" not in cols:
        conn.execute("ALTER TABLE workshop_submissions ADD COLUMN participant_key TEXT")
    if "content_json" not in cols:
        # Structured field values alongside the rendered content_text, so a
        # participant returning to an activity gets their answers back in the
        # right boxes (WCAG 3.3.7 Redundant Entry) rather than a blank form.
        conn.execute("ALTER TABLE workshop_submissions ADD COLUMN content_json TEXT")
    if "is_draft" not in cols:
        conn.execute("ALTER TABLE workshop_submissions ADD COLUMN is_draft INTEGER NOT NULL DEFAULT 0")
    session_cols = [str(r["name"]) for r in conn.execute("PRAGMA table_info(workshop_sessions)").fetchall()]
    if "facilitator_key" not in session_cols:
        conn.execute("ALTER TABLE workshop_sessions ADD COLUMN facilitator_key TEXT")
    # One stored response per participant per activity. Re-saving an activity
    # must revise that row, not append another copy of the same work.
    #
    # Existing deployments predate this rule and may already hold duplicates
    # from repeated saves, which would make the unique index fail to build.
    # Collapse them first, keeping the most recently updated row so no
    # participant loses their latest answer.
    # Peer feedback and follow-through items reference submissions by id, so
    # move anything hanging off a duplicate onto the survivor before deleting.
    # Orphaned feedback would simply vanish from the gallery.
    survivor_map = (
        "SELECT d.id AS dup_id, k.keep_id AS keep_id FROM workshop_submissions d "
        "JOIN ("
        " SELECT session_code, activity_key, participant_key, id AS keep_id FROM ("
        "  SELECT session_code, activity_key, participant_key, id, ROW_NUMBER() OVER ("
        "   PARTITION BY session_code, activity_key, participant_key"
        "   ORDER BY updated_at_utc DESC, id DESC"
        "  ) AS rn FROM workshop_submissions WHERE participant_key IS NOT NULL"
        " ) WHERE rn = 1"
        ") k ON d.session_code = k.session_code AND d.activity_key = k.activity_key "
        "AND d.participant_key = k.participant_key "
        "WHERE d.participant_key IS NOT NULL AND d.id <> k.keep_id"
    )
    remaps = conn.execute(survivor_map).fetchall()
    for row in remaps:
        conn.execute(
            "UPDATE workshop_feedback SET submission_id=? WHERE submission_id=?",
            (int(row["keep_id"]), int(row["dup_id"])),
        )
        conn.execute(
            "UPDATE workshop_follow_through SET source_submission_id=? WHERE source_submission_id=?",
            (int(row["keep_id"]), int(row["dup_id"])),
        )

    conn.execute(
        "DELETE FROM workshop_submissions WHERE participant_key IS NOT NULL AND id NOT IN ("
        " SELECT id FROM ("
        "  SELECT id, ROW_NUMBER() OVER ("
        "   PARTITION BY session_code, activity_key, participant_key"
        "   ORDER BY updated_at_utc DESC, id DESC"
        "  ) AS rn"
        "  FROM workshop_submissions WHERE participant_key IS NOT NULL"
        " ) WHERE rn = 1"
        ")"
    )
    conn.execute(
        "CREATE UNIQUE INDEX IF NOT EXISTS idx_workshop_submissions_unique_owner "
        "ON workshop_submissions(session_code, activity_key, participant_key) "
        "WHERE participant_key IS NOT NULL"
    )
    conn.execute(
        "CREATE INDEX IF NOT EXISTS idx_workshop_submissions_session_activity "
        "ON workshop_submissions(session_code, activity_key)"
    )
    conn.execute(
        "CREATE INDEX IF NOT EXISTS idx_workshop_submissions_participant "
        "ON workshop_submissions(session_code, participant_key, updated_at_utc)"
    )
    conn.execute(
        "CREATE INDEX IF NOT EXISTS idx_workshop_conference_codes_session "
        "ON workshop_conference_codes(session_code, active)"
    )
    conn.commit()


def normalize_session_code(raw: str) -> str:
    value = (raw or "").strip()
    if not _SESSION_CODE_RE.match(value):
        raise ValueError("Session code must be 3-64 chars (letters, numbers, dash, underscore).")
    return value


def ensure_session(session_code: str, *, title: str = "GLOW Workshop Session", event_name: str = "") -> None:
    code = normalize_session_code(session_code)
    conn = _conn()
    conn.execute(
        "INSERT OR IGNORE INTO workshop_sessions (session_code, title, event_name, created_at_utc) VALUES (?, ?, ?, ?)",
        (code, title, event_name, _utc_now()),
    )
    conn.commit()
    conn.close()


def get_session(session_code: str) -> dict | None:
    code = normalize_session_code(session_code)
    conn = _conn()
    row = conn.execute(
        "SELECT session_code, title, event_name, created_at_utc FROM workshop_sessions WHERE session_code=?",
        (code,),
    ).fetchone()
    conn.close()
    return dict(row) if row else None


def save_submission(
    session_code: str,
    activity_key: str,
    display_name: str,
    content_text: str,
    *,
    participant_key: str | None = None,
    anonymity_mode: bool = False,
    field_values: dict[str, str] | None = None,
    is_draft: bool = False,
) -> int:
    """Store a participant's response, replacing their previous one.

    This is an upsert, not an append. A participant who revises an activity is
    revising one answer; the old behaviour left a fresh row behind on every
    save, which inflated the gallery, double-counted progress and made the
    facilitator's completion numbers meaningless by mid-morning.

    Submissions with no ``participant_key`` (seeded or legacy rows) keep the
    old insert behaviour, since there is no owner to key the revision on.
    """
    code = normalize_session_code(session_code)
    now = _utc_now()
    key = (participant_key or "").strip() or None
    name = (display_name or "Participant").strip() or "Participant"
    payload = json.dumps(field_values or {}, ensure_ascii=False)
    conn = _conn()

    if key is None:
        cur = conn.execute(
            "INSERT INTO workshop_submissions "
            "(session_code, activity_key, participant_key, display_name, anonymity_mode, content_text, content_json, is_draft, created_at_utc, updated_at_utc) "
            "VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
            (code, activity_key, None, name, int(bool(anonymity_mode)), content_text, payload, int(bool(is_draft)), now, now),
        )
        conn.commit()
        new_id = int(cur.lastrowid)
        conn.close()
        return new_id

    conn.execute(
        "INSERT INTO workshop_submissions "
        "(session_code, activity_key, participant_key, display_name, anonymity_mode, content_text, content_json, is_draft, created_at_utc, updated_at_utc) "
        "VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?) "
        # The unique index is partial, so the conflict target must repeat its
        # WHERE clause or SQLite will not match it.
        "ON CONFLICT(session_code, activity_key, participant_key) WHERE participant_key IS NOT NULL DO UPDATE SET "
        "display_name=excluded.display_name, anonymity_mode=excluded.anonymity_mode, "
        "content_text=excluded.content_text, content_json=excluded.content_json, "
        # A finished response must never be demoted back to a draft. Someone
        # returning to a completed activity to reword one sentence and pressing
        # "Save draft" would otherwise pull their work out of the gallery, the
        # room count and their own passport -- under a success message.
        "is_draft=CASE WHEN workshop_submissions.is_draft = 0 THEN 0 ELSE excluded.is_draft END, "
        "updated_at_utc=excluded.updated_at_utc",
        (code, activity_key, key, name, int(bool(anonymity_mode)), content_text, payload, int(bool(is_draft)), now, now),
    )
    conn.commit()
    row = conn.execute(
        "SELECT id FROM workshop_submissions WHERE session_code=? AND activity_key=? AND participant_key=?",
        (code, activity_key, key),
    ).fetchone()
    conn.close()
    return int(row["id"]) if row else 0


def get_participant_submission(session_code: str, participant_key: str, activity_key: str) -> dict | None:
    """A participant's stored answer for one activity, if any."""
    code = normalize_session_code(session_code)
    key = (participant_key or "").strip()
    if not key:
        return None
    conn = _conn()
    row = conn.execute(
        "SELECT id, session_code, activity_key, participant_key, display_name, anonymity_mode, content_text, content_json, is_draft, created_at_utc, updated_at_utc "
        "FROM workshop_submissions WHERE session_code=? AND participant_key=? AND activity_key=?",
        (code, key, activity_key),
    ).fetchone()
    conn.close()
    return dict(row) if row else None


def decode_field_values(row: dict | None) -> dict[str, str]:
    """Field values stored with a submission, or an empty mapping."""
    if not row:
        return {}
    raw = (row.get("content_json") or "").strip()
    if not raw:
        return {}
    try:
        decoded = json.loads(raw)
    except (TypeError, ValueError):
        return {}
    if not isinstance(decoded, dict):
        return {}
    return {str(k): str(v) for k, v in decoded.items()}


def _submission_filters(
    code: str,
    activity_key: str | None,
    include_drafts: bool,
) -> tuple[str, list]:
    clause = "session_code=?"
    params: list = [code]
    if activity_key:
        clause += " AND activity_key=?"
        params.append(activity_key)
    if not include_drafts:
        clause += " AND is_draft=0"
    return clause, params


def count_submissions(
    session_code: str,
    *,
    activity_key: str | None = None,
    include_drafts: bool = False,
) -> int:
    code = normalize_session_code(session_code)
    clause, params = _submission_filters(code, activity_key, include_drafts)
    conn = _conn()
    row = conn.execute(
        f"SELECT COUNT(*) AS n FROM workshop_submissions WHERE {clause}",
        params,
    ).fetchone()
    conn.close()
    return int(row["n"]) if row else 0


def list_submissions(
    session_code: str,
    *,
    activity_key: str | None = None,
    include_drafts: bool = False,
    limit: int | None = None,
    offset: int = 0,
) -> list[dict]:
    """Submissions for a session, newest first.

    Drafts are excluded by default: a draft is unfinished work and does not
    belong in the shared gallery, the room-wide counts, or an export.

    ``limit``/``offset`` exist because a full-day session with 25 participants
    produces hundreds of submissions, and the gallery used to render every one
    of them on a single page.
    """
    code = normalize_session_code(session_code)
    clause, params = _submission_filters(code, activity_key, include_drafts)
    sql = f"SELECT id, session_code, activity_key, participant_key, display_name, anonymity_mode, content_text, content_json, is_draft, created_at_utc, updated_at_utc FROM workshop_submissions WHERE {clause} ORDER BY updated_at_utc DESC, id DESC"
    if limit is not None:
        sql += " LIMIT ? OFFSET ?"
        params.extend([int(limit), max(0, int(offset))])
    conn = _conn()
    rows = conn.execute(sql, params).fetchall()
    conn.close()
    return [dict(r) for r in rows]


def list_submissions_for_participant(
    session_code: str,
    participant_key: str,
    *,
    include_drafts: bool = True,
) -> list[dict]:
    """A participant's own submissions.

    Drafts are included here by default -- they are the participant's own
    unfinished work and they need to see it to carry on.
    """
    code = normalize_session_code(session_code)
    key = (participant_key or "").strip()
    if not key:
        return []
    sql = (
        "SELECT id, session_code, activity_key, participant_key, display_name, anonymity_mode, content_text, content_json, is_draft, created_at_utc, updated_at_utc FROM workshop_submissions "
        "WHERE session_code=? AND participant_key=?"
    )
    if not include_drafts:
        sql += " AND is_draft=0"
    sql += " ORDER BY updated_at_utc DESC, id DESC"
    conn = _conn()
    rows = conn.execute(sql, (code, key)).fetchall()
    conn.close()
    return [dict(r) for r in rows]


def create_or_update_participant(
    session_code: str,
    display_name: str,
    *,
    participant_key: str | None = None,
    login_email: str | None = None,
) -> dict:
    code = normalize_session_code(session_code)
    name = (display_name or "Participant").strip() or "Participant"
    key = (participant_key or "").strip() or secrets.token_urlsafe(24)
    now = _utc_now()
    email = (login_email or "").strip() or None
    conn = _conn()
    row = conn.execute(
        "SELECT participant_key, session_code, display_name, login_email, created_at_utc, updated_at_utc, last_seen_at_utc "
        "FROM workshop_participants WHERE participant_key=?",
        (key,),
    ).fetchone()
    if row and str(row["session_code"]) == code:
        conn.execute(
            "UPDATE workshop_participants SET display_name=?, login_email=COALESCE(?, login_email), updated_at_utc=?, last_seen_at_utc=? "
            "WHERE participant_key=?",
            (name, email, now, now, key),
        )
    elif row:
        # Prevent cross-session token reuse; issue new token.
        key = secrets.token_urlsafe(24)
        conn.execute(
            "INSERT INTO workshop_participants (participant_key, session_code, display_name, login_email, created_at_utc, updated_at_utc, last_seen_at_utc) "
            "VALUES (?, ?, ?, ?, ?, ?, ?)",
            (key, code, name, email, now, now, now),
        )
    else:
        conn.execute(
            "INSERT INTO workshop_participants (participant_key, session_code, display_name, login_email, created_at_utc, updated_at_utc, last_seen_at_utc) "
            "VALUES (?, ?, ?, ?, ?, ?, ?)",
            (key, code, name, email, now, now, now),
        )
    conn.commit()
    out = conn.execute(
        "SELECT participant_key, session_code, display_name, login_email, created_at_utc, updated_at_utc, last_seen_at_utc "
        "FROM workshop_participants WHERE participant_key=?",
        (key,),
    ).fetchone()
    conn.close()
    return dict(out) if out else {"participant_key": key, "session_code": code, "display_name": name}


def get_submission(session_code: str, submission_id: int) -> dict | None:
    """One submission, scoped to its session so an id cannot cross sessions."""
    code = normalize_session_code(session_code)
    conn = _conn()
    row = conn.execute(
        "SELECT id, session_code, activity_key, participant_key, display_name, anonymity_mode, "
        "content_text, content_json, is_draft, created_at_utc, updated_at_utc "
        "FROM workshop_submissions WHERE session_code=? AND id=?",
        (code, int(submission_id)),
    ).fetchone()
    conn.close()
    return dict(row) if row else None


def participants_due_for_nudge(session_code: str, *, days: int = 30) -> list[dict]:
    """Participants whose 30-day commitment has come due.

    A participant qualifies when they left an email address, finished the
    30-Day Action Plan, saved it at least *days* ago, and have not been
    nudged before. Nobody is emailed twice, and nobody who never gave an
    address is emailed at all.
    """
    code = normalize_session_code(session_code)
    cutoff = (datetime.now(UTC) - timedelta(days=int(days))).isoformat()
    conn = _conn()
    rows = conn.execute(
        "SELECT p.participant_key, p.display_name, p.login_email, "
        "       s.content_json, s.updated_at_utc "
        "FROM workshop_participants p "
        "JOIN workshop_submissions s "
        "  ON s.participant_key = p.participant_key AND s.session_code = p.session_code "
        "LEFT JOIN workshop_nudges n "
        "  ON n.participant_key = p.participant_key AND n.session_code = p.session_code "
        "WHERE p.session_code = ? "
        "  AND s.activity_key = 'action_plan_30_day' "
        "  AND COALESCE(s.is_draft, 0) = 0 "
        "  AND s.updated_at_utc <= ? "
        "  AND p.login_email IS NOT NULL AND TRIM(p.login_email) <> '' "
        "  AND n.participant_key IS NULL "
        "ORDER BY s.updated_at_utc",
        (code, cutoff),
    ).fetchall()
    conn.close()
    return [dict(row) for row in rows]


def record_nudge(session_code: str, participant_key: str) -> None:
    """Remember that this person has been asked, so they are asked once."""
    code = normalize_session_code(session_code)
    conn = _conn()
    conn.execute(
        "INSERT OR REPLACE INTO workshop_nudges (session_code, participant_key, sent_at_utc) "
        "VALUES (?, ?, ?)",
        (code, (participant_key or "").strip(), _utc_now()),
    )
    conn.commit()
    conn.close()


def count_participants(session_code: str) -> int:
    """How many people have joined this session."""
    code = normalize_session_code(session_code)
    conn = _conn()
    row = conn.execute(
        "SELECT COUNT(*) AS total FROM workshop_participants WHERE session_code=?",
        (code,),
    ).fetchone()
    conn.close()
    return int(row["total"]) if row else 0


def get_participant(participant_key: str) -> dict | None:
    key = (participant_key or "").strip()
    if not key:
        return None
    conn = _conn()
    row = conn.execute(
        "SELECT participant_key, session_code, display_name, login_email, created_at_utc, updated_at_utc, last_seen_at_utc "
        "FROM workshop_participants WHERE participant_key=?",
        (key,),
    ).fetchone()
    if row:
        conn.execute(
            "UPDATE workshop_participants SET last_seen_at_utc=?, updated_at_utc=? WHERE participant_key=?",
            (_utc_now(), _utc_now(), key),
        )
        conn.commit()
    conn.close()
    return dict(row) if row else None


# ---------------------------------------------------------------------------
# Return links
# ---------------------------------------------------------------------------
#
# Participant identity is a cookie on one device. Close the laptop, switch to
# a phone, or let the browser clear cookies, and a participant's whole day is
# unreachable -- and the 30-day action plan, which is the point of the
# workshop, has no delivery vehicle. A return link restores the same
# participant_key on any device from an emailed, single-use, expiring token.
#
# The token is a bearer credential for one person's workshop work, so only its
# SHA-256 hash is stored: a copy of the database does not yield working links.

RETURN_LINK_TTL_DAYS = 45


def return_link_ttl_days() -> int:
    """Days a return link stays valid.

    The default outlives the 30-day action plan, so the follow-up nudge can
    link a participant straight back into their own commitments.
    """
    raw = (os.environ.get("GLOW_WORKSHOP_RETURN_LINK_TTL_DAYS") or "").strip()
    if raw.isdigit() and int(raw) > 0:
        return int(raw)
    return RETURN_LINK_TTL_DAYS


def _hash_return_token(token: str) -> str:
    return hashlib.sha256(token.encode("utf-8")).hexdigest()


def create_return_link(session_code: str, participant_key: str, *, ttl_days: int | None = None) -> str:
    """Issue a single-use return token and give back the raw value to email.

    The raw token is returned once and never stored; only its hash is kept.
    """
    code = normalize_session_code(session_code)
    key = (participant_key or "").strip()
    if not key:
        raise ValueError("A participant key is required.")
    days = ttl_days if ttl_days and ttl_days > 0 else return_link_ttl_days()
    token = secrets.token_urlsafe(32)
    now = datetime.now(UTC)
    conn = _conn()
    conn.execute(
        "INSERT INTO workshop_return_links "
        "(token_hash, session_code, participant_key, created_at_utc, expires_at_utc, used_at_utc) "
        "VALUES (?, ?, ?, ?, ?, NULL)",
        (
            _hash_return_token(token),
            code,
            key,
            now.isoformat(),
            (now + timedelta(days=days)).isoformat(),
        ),
    )
    conn.commit()
    conn.close()
    return token


def consume_return_link(token: str) -> tuple[str, dict | None]:
    """Spend a return token.

    Returns ``(status, participant)`` where status is one of ``ok``,
    ``invalid``, ``used``, ``expired`` or ``gone``. The caller needs the
    distinction: "this link was already used" and "this link expired" are
    different instructions for the person holding it, and neither should be
    reported as a generic failure.
    """
    raw = (token or "").strip()
    if not raw:
        return "invalid", None

    conn = _conn()
    row = conn.execute(
        "SELECT token_hash, session_code, participant_key, expires_at_utc, used_at_utc "
        "FROM workshop_return_links WHERE token_hash=?",
        (_hash_return_token(raw),),
    ).fetchone()
    if row is None:
        conn.close()
        return "invalid", None
    if row["used_at_utc"]:
        conn.close()
        return "used", None
    try:
        expires = datetime.fromisoformat(str(row["expires_at_utc"]))
    except ValueError:
        expires = datetime.now(UTC) - timedelta(seconds=1)
    if expires.tzinfo is None:
        expires = expires.replace(tzinfo=UTC)
    if expires < datetime.now(UTC):
        conn.close()
        return "expired", None

    participant = conn.execute(
        "SELECT participant_key, session_code, display_name, login_email, created_at_utc, updated_at_utc, last_seen_at_utc "
        "FROM workshop_participants WHERE participant_key=?",
        (str(row["participant_key"]),),
    ).fetchone()
    if participant is None:
        # The participant record was removed after the link was issued.
        conn.close()
        return "gone", None

    now = _utc_now()
    conn.execute(
        "UPDATE workshop_return_links SET used_at_utc=? WHERE token_hash=?",
        (now, str(row["token_hash"])),
    )
    conn.execute(
        "UPDATE workshop_participants SET last_seen_at_utc=?, updated_at_utc=? WHERE participant_key=?",
        (now, now, str(row["participant_key"])),
    )
    conn.commit()
    conn.close()
    return "ok", dict(participant)


def bind_participant_login(participant_key: str, login_email: str) -> None:
    key = (participant_key or "").strip()
    email = (login_email or "").strip()
    if not key or not email:
        return
    conn = _conn()
    conn.execute(
        "UPDATE workshop_participants SET login_email=?, updated_at_utc=?, last_seen_at_utc=? WHERE participant_key=?",
        (email, _utc_now(), _utc_now(), key),
    )
    conn.commit()
    conn.close()


def upsert_conference_code(
    access_code: str,
    *,
    session_code: str,
    session_title: str,
    event_name: str = "",
    active: bool = True,
    facilitator_key: str = "",
) -> None:
    access = (access_code or "").strip().upper()
    if not access:
        raise ValueError("Conference access code is required.")
    code = normalize_session_code(session_code)
    title = (session_title or "GLOW Workshop Session").strip() or "GLOW Workshop Session"
    event = (event_name or "").strip()
    now = _utc_now()
    ensure_session(code, title=title, event_name=event)
    if (facilitator_key or "").strip():
        set_facilitator_key(code, facilitator_key)
    conn = _conn()
    conn.execute(
        "INSERT INTO workshop_conference_codes (access_code, session_code, session_title, event_name, active, created_at_utc, updated_at_utc) "
        "VALUES (?, ?, ?, ?, ?, ?, ?) "
        "ON CONFLICT(access_code) DO UPDATE SET "
        "session_code=excluded.session_code, session_title=excluded.session_title, event_name=excluded.event_name, active=excluded.active, updated_at_utc=excluded.updated_at_utc",
        (access, code, title, event, int(bool(active)), now, now),
    )
    conn.commit()
    conn.close()


def get_conference_code(access_code: str) -> dict | None:
    access = (access_code or "").strip().upper()
    if not access:
        return None
    conn = _conn()
    row = conn.execute(
        "SELECT access_code, session_code, session_title, event_name, active, created_at_utc, updated_at_utc "
        "FROM workshop_conference_codes WHERE access_code=? AND active=1",
        (access,),
    ).fetchone()
    conn.close()
    return dict(row) if row else None


def load_conference_codes_from_file() -> int:
    """Load conference access code mappings from instance/workshop_conference_codes.json."""
    p = Path(current_app.instance_path) / "workshop_conference_codes.json"
    if not p.exists():
        return 0
    try:
        payload = json.loads(p.read_text(encoding="utf-8"))
    except Exception:
        return 0
    if not isinstance(payload, list):
        return 0

    count = 0
    for item in payload:
        if not isinstance(item, dict):
            continue
        access = (str(item.get("access_code", "")) or "").strip()
        session_code = (str(item.get("session_code", "")) or "").strip()
        if not access or not session_code:
            continue
        try:
            upsert_conference_code(
                access,
                session_code=session_code,
                session_title=(str(item.get("session_title", "GLOW Workshop Session")) or "").strip() or "GLOW Workshop Session",
                event_name=(str(item.get("event_name", "")) or "").strip(),
                active=bool(item.get("active", True)),
                facilitator_key=(str(item.get("facilitator_key", "")) or "").strip(),
            )
            count += 1
        except Exception:
            continue
    return count


def load_conference_codes_from_env() -> int:
    """Load mappings from WORKSHOP_CONFERENCE_CODES_JSON env var (JSON list)."""
    raw = os.environ.get("WORKSHOP_CONFERENCE_CODES_JSON", "").strip()
    if not raw:
        return 0
    try:
        payload = json.loads(raw)
    except Exception:
        return 0
    if not isinstance(payload, list):
        return 0
    count = 0
    for item in payload:
        if not isinstance(item, dict):
            continue
        access = (str(item.get("access_code", "")) or "").strip()
        session_code = (str(item.get("session_code", "")) or "").strip()
        if not access or not session_code:
            continue
        try:
            upsert_conference_code(
                access,
                session_code=session_code,
                session_title=(str(item.get("session_title", "GLOW Workshop Session")) or "").strip() or "GLOW Workshop Session",
                event_name=(str(item.get("event_name", "")) or "").strip(),
                active=bool(item.get("active", True)),
                facilitator_key=(str(item.get("facilitator_key", "")) or "").strip(),
            )
            count += 1
        except Exception:
            continue
    return count


def add_feedback(
    session_code: str,
    submission_id: int,
    reviewer_display_name: str,
    strength: str,
    risk_or_gap: str,
    recommended_safeguard: str,
    reuse_suggestion: str,
) -> int:
    code = normalize_session_code(session_code)
    conn = _conn()
    cur = conn.execute(
        "INSERT INTO workshop_feedback (session_code, submission_id, reviewer_display_name, strength, risk_or_gap, recommended_safeguard, reuse_suggestion, created_at_utc) "
        "VALUES (?, ?, ?, ?, ?, ?, ?, ?)",
        (
            code,
            int(submission_id),
            (reviewer_display_name or "Peer Reviewer").strip() or "Peer Reviewer",
            strength,
            risk_or_gap,
            recommended_safeguard,
            reuse_suggestion,
            _utc_now(),
        ),
    )
    conn.commit()
    feedback_id = int(cur.lastrowid)
    conn.close()
    return feedback_id


def save_follow_through_item(
    session_code: str,
    kind: str,
    title: str,
    details: str,
    owner_name: str,
    *,
    due_date: str | None = None,
    source_submission_id: int | None = None,
) -> int:
    code = normalize_session_code(session_code)
    now = _utc_now()
    conn = _conn()
    cur = conn.execute(
        "INSERT INTO workshop_follow_through (session_code, source_submission_id, kind, title, details, owner_name, due_date, status, created_at_utc, updated_at_utc) "
        "VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
        (
            code,
            None if source_submission_id is None else int(source_submission_id),
            (kind or "action_commitment").strip() or "action_commitment",
            (title or "Workshop follow-through item").strip() or "Workshop follow-through item",
            (details or "").strip(),
            (owner_name or "Participant").strip() or "Participant",
            (due_date or "").strip() or None,
            "open",
            now,
            now,
        ),
    )
    conn.commit()
    new_id = int(cur.lastrowid)
    conn.close()
    return new_id


def list_follow_through_items(session_code: str, *, kind: str | None = None) -> list[dict]:
    code = normalize_session_code(session_code)
    conn = _conn()
    if kind:
        rows = conn.execute(
            "SELECT id, session_code, source_submission_id, kind, title, details, owner_name, due_date, status, created_at_utc, updated_at_utc "
            "FROM workshop_follow_through WHERE session_code=? AND kind=? ORDER BY updated_at_utc DESC, id DESC",
            (code, kind),
        ).fetchall()
    else:
        rows = conn.execute(
            "SELECT id, session_code, source_submission_id, kind, title, details, owner_name, due_date, status, created_at_utc, updated_at_utc "
            "FROM workshop_follow_through WHERE session_code=? ORDER BY updated_at_utc DESC, id DESC",
            (code,),
        ).fetchall()
    conn.close()
    return [dict(r) for r in rows]


def update_follow_through_status(session_code: str, item_id: int, status: str) -> None:
    code = normalize_session_code(session_code)
    normalized = (status or "").strip().lower()
    if normalized not in {"open", "done", "paused"}:
        raise ValueError("Status must be open, done, or paused.")
    conn = _conn()
    conn.execute(
        "UPDATE workshop_follow_through SET status=?, updated_at_utc=? WHERE session_code=? AND id=?",
        (normalized, _utc_now(), code, int(item_id)),
    )
    conn.commit()
    conn.close()


def list_feedback_for_session(session_code: str) -> dict[int, list[dict]]:
    code = normalize_session_code(session_code)
    conn = _conn()
    rows = conn.execute(
        "SELECT id, submission_id, reviewer_display_name, strength, risk_or_gap, recommended_safeguard, reuse_suggestion, created_at_utc "
        "FROM workshop_feedback WHERE session_code=? ORDER BY created_at_utc DESC",
        (code,),
    ).fetchall()
    conn.close()

    grouped: dict[int, list[dict]] = {}
    for r in rows:
        item = dict(r)
        grouped.setdefault(int(item["submission_id"]), []).append(item)
    return grouped


def redact_submissions_for_export(rows: list[dict]) -> list[dict]:
    """Strip participant-identifying material from exported submissions.

    Two things must never leave this function intact:

    * ``participant_key`` is the literal value of the ``glow_workshop_participant``
      cookie, and :func:`get_participant` accepts it as a bearer credential with
      no further check. Exporting it hands anyone who downloads the file the
      ability to assume that participant's workshop identity.
    * ``display_name`` must be masked whenever the participant ticked the
      anonymity box. The Markdown, HTML and DOCX exports already honour that
      choice; the JSON export must not become the way around it.

    ``login_email`` is dropped for the same reason even though the current
    submission query does not select it -- so a future column addition cannot
    silently start leaking it.
    """
    redacted: list[dict] = []
    for row in rows:
        item = dict(row)
        item.pop("participant_key", None)
        item.pop("login_email", None)
        if int(item.get("anonymity_mode", 0) or 0):
            item["display_name"] = "Anonymous participant"
        redacted.append(item)
    return redacted


def set_facilitator_key(session_code: str, facilitator_key: str) -> None:
    code = normalize_session_code(session_code)
    value = (facilitator_key or "").strip() or None
    conn = _conn()
    conn.execute(
        "UPDATE workshop_sessions SET facilitator_key=? WHERE session_code=?",
        (value, code),
    )
    conn.commit()
    conn.close()


def get_facilitator_key(session_code: str) -> str | None:
    code = normalize_session_code(session_code)
    conn = _conn()
    row = conn.execute(
        "SELECT facilitator_key FROM workshop_sessions WHERE session_code=?",
        (code,),
    ).fetchone()
    conn.close()
    if not row:
        return None
    value = (row["facilitator_key"] or "").strip()
    return value or None


def export_session_markdown(session_code: str, *, session_title: str = "GLOW Workshop Session") -> str:
    code = normalize_session_code(session_code)
    submissions = list_submissions(code)
    feedback = list_feedback_for_session(code)

    lines: list[str] = []
    lines.append(f"# {session_title}")
    lines.append("")
    lines.append(f"Session code: `{code}`")
    lines.append("")
    lines.append("## Submissions")
    lines.append("")

    if not submissions:
        lines.append("No submissions captured yet.")
        lines.append("")
        return "\n".join(lines)

    for s in submissions:
        submitter = "Anonymous participant" if int(s.get("anonymity_mode", 0)) else s.get("display_name", "Participant")
        lines.append(f"### {s.get('activity_key', 'unknown_activity')}")
        lines.append("")
        lines.append(f"- Submitter: {submitter}")
        lines.append(f"- Updated (UTC): {s.get('updated_at_utc', '')}")
        lines.append("")
        lines.append("```text")
        lines.append((s.get("content_text") or "").strip())
        lines.append("```")
        lines.append("")

        items = feedback.get(int(s.get("id", 0)), [])
        if items:
            lines.append("#### Peer Feedback")
            lines.append("")
            for f in items:
                lines.append(f"- Reviewer: {f.get('reviewer_display_name', 'Peer Reviewer')}")
                lines.append(f"  - Strength: {f.get('strength', '')}")
                lines.append(f"  - Risk or gap: {f.get('risk_or_gap', '')}")
                lines.append(f"  - Recommended safeguard: {f.get('recommended_safeguard', '')}")
                lines.append(f"  - Reuse suggestion: {f.get('reuse_suggestion', '')}")
            lines.append("")

    return "\n".join(lines)


def export_session_json(session_code: str, *, session_title: str = "GLOW Workshop Session") -> str:
    code = normalize_session_code(session_code)
    payload = {
        "session": get_session(code) or {"session_code": code, "title": session_title},
        "submissions": redact_submissions_for_export(list_submissions(code)),
        "feedback_by_submission": list_feedback_for_session(code),
    }
    return json.dumps(payload, indent=2)


def export_session_html(session_code: str, *, session_title: str = "GLOW Workshop Session") -> str:
    code = normalize_session_code(session_code)
    submissions = list_submissions(code)
    feedback = list_feedback_for_session(code)

    cards: list[str] = []
    if not submissions:
        cards.append("<p>No submissions captured yet.</p>")
    else:
        for s in submissions:
            submitter = "Anonymous participant" if int(s.get("anonymity_mode", 0)) else escape(s.get("display_name", "Participant"))
            body = escape((s.get("content_text") or "").strip())
            parts = [
                f"<article><h2>{escape(s.get('activity_key', 'unknown_activity'))}</h2>",
                f"<p><strong>Submitter:</strong> {submitter}</p>",
                f"<p><strong>Updated (UTC):</strong> {escape(s.get('updated_at_utc', ''))}</p>",
                f"<pre>{body}</pre>",
            ]
            items = feedback.get(int(s.get("id", 0)), [])
            if items:
                parts.append("<h3>Peer Feedback</h3><ul>")
                for f in items:
                    parts.append(
                        "<li>"
                        f"<strong>{escape(f.get('reviewer_display_name', 'Peer Reviewer'))}</strong>: "
                        f"{escape(f.get('strength', ''))} "
                        f"(Risk/gap: {escape(f.get('risk_or_gap', ''))}; "
                        f"Safeguard: {escape(f.get('recommended_safeguard', ''))}; "
                        f"Reuse: {escape(f.get('reuse_suggestion', ''))})"
                        "</li>"
                    )
                parts.append("</ul>")
            parts.append("</article>")
            cards.append("\n".join(parts))

    title = escape(session_title)
    return f"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>{title}</title>
  <style>
    body {{ font-family: system-ui, sans-serif; line-height: 1.5; margin: 0; background: #f8fafc; }}
    main {{ max-width: 72rem; margin: 0 auto; padding: 1rem; }}
    article {{ background: #fff; border: 1px solid #ccd; border-radius: .5rem; padding: .875rem; margin: .75rem 0; }}
    pre {{ white-space: pre-wrap; }}
  </style>
</head>
<body>
  <main>
    <h1>{title}</h1>
    <p><strong>Session code:</strong> {escape(code)}</p>
    {''.join(cards)}
  </main>
</body>
</html>"""


def export_session_docx_bytes(session_code: str, *, session_title: str = "GLOW Workshop Session") -> bytes:
    code = normalize_session_code(session_code)
    submissions = list_submissions(code)
    feedback = list_feedback_for_session(code)

    # python-docx is available via desktop package dependency in this repository.
    from docx import Document  # type: ignore

    doc = Document()
    doc.add_heading(session_title, level=1)
    doc.add_paragraph(f"Session code: {code}")

    if not submissions:
        doc.add_paragraph("No submissions captured yet.")
    else:
        for s in submissions:
            activity = s.get("activity_key", "unknown_activity")
            submitter = "Anonymous participant" if int(s.get("anonymity_mode", 0)) else s.get("display_name", "Participant")
            doc.add_heading(str(activity), level=2)
            doc.add_paragraph(f"Submitter: {submitter}")
            doc.add_paragraph(f"Updated (UTC): {s.get('updated_at_utc', '')}")
            doc.add_paragraph((s.get("content_text") or "").strip() or "(empty)")

            items = feedback.get(int(s.get("id", 0)), [])
            if items:
                doc.add_heading("Peer Feedback", level=3)
                for f in items:
                    doc.add_paragraph(
                        f"Reviewer: {f.get('reviewer_display_name', 'Peer Reviewer')}",
                        style="List Bullet",
                    )
                    doc.add_paragraph(f"Strength: {f.get('strength', '')}")
                    doc.add_paragraph(f"Risk/gap: {f.get('risk_or_gap', '')}")
                    doc.add_paragraph(f"Safeguard: {f.get('recommended_safeguard', '')}")
                    doc.add_paragraph(f"Reuse suggestion: {f.get('reuse_suggestion', '')}")

    out = BytesIO()
    doc.save(out)
    return out.getvalue()


def export_follow_through_markdown(session_code: str, *, session_title: str = "GLOW Workshop Session") -> str:
    code = normalize_session_code(session_code)
    items = list_follow_through_items(code)

    lines: list[str] = []
    lines.append(f"# {session_title} Follow-Through")
    lines.append("")
    lines.append(f"Session code: `{code}`")
    lines.append("")
    lines.append("## Saved coaching artifacts")
    lines.append("")

    if not items:
        lines.append("No follow-through items saved yet.")
        lines.append("")
        return "\n".join(lines)

    for item in items:
        lines.append(f"### {item.get('title', 'Follow-through item')}")
        lines.append("")
        lines.append(f"- Kind: {item.get('kind', 'action_commitment')}")
        lines.append(f"- Owner: {item.get('owner_name', 'Participant')}")
        if item.get("due_date"):
            lines.append(f"- Due date: {item.get('due_date')}")
        lines.append(f"- Status: {item.get('status', 'open')}")
        lines.append("")
        lines.append("```text")
        lines.append(item.get("details", ""))
        lines.append("```")
        lines.append("")

    return "\n".join(lines)
