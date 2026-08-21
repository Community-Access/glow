"""Blank worksheet packs, in HTML and Word.

The workshop promises that the day works without the web app: on paper, at a
table, for someone whose institution blocks the tool, whose wifi has failed,
or who simply prefers a pen. Exports covered *completed* work; there was
nothing to hand someone who wants to start.

Both formats are generated from the same activity definitions the app renders,
passed in by the caller, so a worksheet cannot drift away from the live form.
The HTML pack is a standalone document rather than an app page: it is
downloaded, printed, or opened offline, so it carries its own styles.

Typography follows the ACB large-print guidance the rest of GLOW enforces --
Arial, 18pt body, generous leading, flush left, no hyphenation. A worksheet
handed out at an accessibility workshop should not need remediation.
"""

from __future__ import annotations

from dataclasses import dataclass
from html import escape
from io import BytesIO

# Writing space. Print needs physical room; the Word pack uses ruled lines so
# there is somewhere to write rather than an ambiguous gap.
_LINES_PER_ROW = 2
_MIN_LINES = 3
_RULE = "_" * 60


@dataclass(frozen=True)
class WorksheetField:
    label: str
    rows: int = 3


@dataclass(frozen=True)
class Worksheet:
    key: str
    title: str
    time: str
    badge: str
    prompt: str
    fields: tuple[WorksheetField, ...]
    scenarios: tuple[tuple[str, str], ...] = ()


def _line_count(field: WorksheetField) -> int:
    return max(_MIN_LINES, int(field.rows or 0) * _LINES_PER_ROW)


_HTML_STYLES = """
    :root { color-scheme: light; }
    body {
      font-family: Arial, Helvetica, sans-serif;
      font-size: 18pt;
      line-height: 1.5;
      letter-spacing: 0.12em;
      word-spacing: 0.16em;
      color: #1a1a1a;
      background: #ffffff;
      text-align: left;
      hyphens: none;
      max-width: 42em;
      margin: 0 auto;
      padding: 1in 0.75in;
    }
    h1 { font-size: 26pt; }
    h2 { font-size: 22pt; margin-top: 2rem; }
    h3 { font-size: 19pt; }
    .meta { font-size: 16pt; }
    .prompt { border-left: 6px solid #1a1a1a; padding-left: 1rem; }
    .write-space {
      border: 2px solid #1a1a1a;
      min-height: 3em;
      margin: 0.5rem 0 1.5rem;
      padding: 0.5rem;
    }
    ol.activities > li { margin-bottom: 0.5rem; }
    @media print {
      body { padding: 0.5in; }
      h2 { page-break-before: always; }
      h2:first-of-type { page-break-before: avoid; }
      .write-space { min-height: 4em; }
    }
"""


def build_worksheet_html(worksheets: list[Worksheet], *, event_name: str = "") -> str:
    """A standalone, printable, screen-reader-friendly worksheet pack."""
    title = "GLOW Workshop Worksheets"
    if event_name.strip():
        title = f"{title} — {event_name.strip()}"

    parts: list[str] = [
        "<!DOCTYPE html>",
        '<html lang="en">',
        "<head>",
        '<meta charset="utf-8">',
        '<meta name="viewport" content="width=device-width, initial-scale=1">',
        f"<title>{escape(title)}</title>",
        f"<style>{_HTML_STYLES}</style>",
        "</head>",
        "<body>",
        f"<h1>{escape(title)}</h1>",
        "<p>",
        "Every activity in the workshop, with space to write. Use these on "
        "paper, in a word processor, or with a screen reader. Nothing here "
        "needs the web app, an account, or an internet connection.",
        "</p>",
        "<h2>Contents</h2>",
        '<ol class="activities">',
    ]
    for sheet in worksheets:
        parts.append(f"<li>{escape(sheet.title)} ({escape(sheet.time)})</li>")
    parts.append("</ol>")

    for sheet in worksheets:
        parts.append(f'<h2 id="{escape(sheet.key)}">{escape(sheet.title)}</h2>')
        parts.append(
            f'<p class="meta">Suggested length: {escape(sheet.time)}. '
            f"Badge: {escape(sheet.badge)}.</p>"
        )
        parts.append(f'<p class="prompt">{escape(sheet.prompt)}</p>')

        if sheet.scenarios:
            parts.append("<h3>Scenarios you can work from</h3>")
            parts.append("<ul>")
            for scenario_title, sector in sheet.scenarios:
                parts.append(
                    f"<li>{escape(scenario_title)} — {escape(sector)}</li>"
                )
            parts.append("</ul>")
            parts.append(
                "<p>Or use a real document of your own. Your own work is the "
                "better choice when you have it.</p>"
            )

        for field in sheet.fields:
            field_id = f"{sheet.key}-{_slug(field.label)}"
            parts.append(f'<h3 id="{escape(field_id)}">{escape(field.label)}</h3>')
            parts.append(
                f'<div class="write-space" role="note" '
                f'aria-label="Writing space for: {escape(field.label)}"></div>'
            )

    parts.append("</body>")
    parts.append("</html>")
    return "\n".join(parts)


def build_worksheet_docx_bytes(worksheets: list[Worksheet], *, event_name: str = "") -> bytes:
    """The same pack as a Word document, in ACB large print."""
    from docx import Document  # type: ignore
    from docx.shared import Pt  # type: ignore

    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(18)
    normal.paragraph_format.space_after = Pt(12)

    title = "GLOW Workshop Worksheets"
    if event_name.strip():
        title = f"{title} - {event_name.strip()}"
    doc.add_heading(title, level=1)
    doc.add_paragraph(
        "Every activity in the workshop, with space to write. Use these on "
        "paper, in a word processor, or with a screen reader. Nothing here "
        "needs the web app, an account, or an internet connection."
    )

    doc.add_heading("Contents", level=2)
    for sheet in worksheets:
        doc.add_paragraph(f"{sheet.title} ({sheet.time})", style="List Number")

    for sheet in worksheets:
        doc.add_page_break()
        doc.add_heading(sheet.title, level=2)
        doc.add_paragraph(f"Suggested length: {sheet.time}. Badge: {sheet.badge}.")
        doc.add_paragraph(sheet.prompt)

        if sheet.scenarios:
            doc.add_heading("Scenarios you can work from", level=3)
            for scenario_title, sector in sheet.scenarios:
                doc.add_paragraph(f"{scenario_title} - {sector}", style="List Bullet")
            doc.add_paragraph(
                "Or use a real document of your own. Your own work is the "
                "better choice when you have it."
            )

        for field in sheet.fields:
            doc.add_heading(field.label, level=3)
            for _ in range(_line_count(field)):
                doc.add_paragraph(_RULE)

    out = BytesIO()
    doc.save(out)
    return out.getvalue()


def _slug(value: str) -> str:
    keep = [c.lower() if c.isalnum() else "-" for c in (value or "").strip()]
    slug = "".join(keep).strip("-")
    while "--" in slug:
        slug = slug.replace("--", "-")
    return slug[:48] or "field"
