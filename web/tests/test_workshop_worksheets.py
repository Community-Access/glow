"""The blank worksheet pack -- the workshop's Tier 0 path.

The day has to work for someone whose institution blocks the tool, whose
wifi has failed, or who would simply rather use a pen. That means a real
worksheet pack, generated from the same activity definitions as the on-screen
form so the two cannot drift apart.
"""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

import pytest
from flask import Flask

from acb_large_print_web.app import create_app
from acb_large_print_web.routes.workshop import (
    ACTIVITY_FIELDS,
    ACTIVITY_ORDER,
    ACTIVITY_PROMPTS,
    _worksheet_pack,
)
from acb_large_print_web.workshop_worksheets import (
    build_worksheet_docx_bytes,
    build_worksheet_html,
)

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


@pytest.fixture()
def app(tmp_path: Path) -> Flask:
    application = create_app({"TESTING": True, "WTF_CSRF_ENABLED": False})
    application.instance_path = str(tmp_path / "instance")
    Path(application.instance_path).mkdir(parents=True, exist_ok=True)
    return application


@pytest.fixture()
def client(app: Flask):
    return app.test_client()


def _pack(app: Flask):
    with app.test_request_context():
        return _worksheet_pack()


# ---------------------------------------------------------------------------
# The pack is the live activity set, not a copy of it
# ---------------------------------------------------------------------------


def test_the_pack_covers_every_activity_in_order(app: Flask):
    pack = _pack(app)

    assert [sheet.key for sheet in pack] == ACTIVITY_ORDER


def test_every_field_on_screen_has_writing_space_on_paper(app: Flask):
    for sheet in _pack(app):
        expected = [
            str(field.get("label", field.get("name", "")))
            for field in ACTIVITY_FIELDS.get(sheet.key, [])
        ]
        if expected:
            assert [field.label for field in sheet.fields] == expected, sheet.key


def test_each_worksheet_carries_its_prompt(app: Flask):
    for sheet in _pack(app):
        assert sheet.prompt == ACTIVITY_PROMPTS[sheet.key]


def test_labs_list_their_scenarios_so_paper_users_get_the_bank_too(app: Flask):
    labs = [sheet for sheet in _pack(app) if sheet.key.startswith("lab_")]

    assert labs
    for lab in labs:
        assert len(lab.scenarios) >= 3, lab.key


# ---------------------------------------------------------------------------
# HTML pack
# ---------------------------------------------------------------------------


def test_html_pack_is_a_complete_accessible_document(app: Flask):
    html = build_worksheet_html(_pack(app))

    assert html.startswith("<!DOCTYPE html>")
    assert '<html lang="en">' in html
    assert html.count("<h1>") == 1
    for sheet in _pack(app):
        assert f'id="{sheet.key}"' in html


def test_html_writing_spaces_are_labelled_for_screen_readers(app: Flask):
    html = build_worksheet_html(_pack(app))

    assert 'aria-label="Writing space for: ' in html
    # One labelled space per field across the whole pack.
    total_fields = sum(len(sheet.fields) for sheet in _pack(app))
    assert html.count("Writing space for:") == total_fields


def test_html_pack_uses_the_acb_type_scale(app: Flask):
    html = build_worksheet_html(_pack(app))

    assert "font-size: 18pt" in html
    assert "font-family: Arial" in html
    assert "hyphens: none" in html


def test_html_pack_names_the_event_when_given_one(app: Flask):
    html = build_worksheet_html(_pack(app), event_name="Accessing Higher Ground")

    assert "Accessing Higher Ground" in html


def test_html_pack_escapes_content_rather_than_trusting_it():
    from acb_large_print_web.workshop_worksheets import Worksheet, WorksheetField

    sheet = Worksheet(
        key="k",
        title="<script>alert(1)</script>",
        time="5 minutes",
        badge="Badge",
        prompt="Prompt",
        fields=(WorksheetField(label="Field", rows=2),),
    )

    html = build_worksheet_html([sheet])

    assert "<script>alert(1)</script>" not in html
    assert "&lt;script&gt;" in html


# ---------------------------------------------------------------------------
# Word pack
# ---------------------------------------------------------------------------


def test_docx_pack_opens_and_carries_every_activity(app: Flask):
    from docx import Document

    payload = build_worksheet_docx_bytes(_pack(app))
    doc = Document(BytesIO(payload))
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)

    for sheet in _pack(app):
        assert sheet.title in text


def test_docx_pack_is_large_print_by_default(app: Flask):
    from docx import Document
    from docx.shared import Pt

    doc = Document(BytesIO(build_worksheet_docx_bytes(_pack(app))))
    normal = doc.styles["Normal"]

    assert normal.font.name == "Arial"
    assert normal.font.size == Pt(18)


def test_docx_pack_leaves_room_to_write(app: Flask):
    from docx import Document

    doc = Document(BytesIO(build_worksheet_docx_bytes(_pack(app))))
    rules = [p for p in doc.paragraphs if p.text.startswith("____")]

    total_fields = sum(len(sheet.fields) for sheet in _pack(app))
    assert len(rules) >= total_fields * 3


# ---------------------------------------------------------------------------
# Routes
# ---------------------------------------------------------------------------


def test_both_packs_download_without_a_session_or_an_account(client):
    html = client.get("/workshop/worksheets.html")
    docx = client.get("/workshop/worksheets.docx")

    assert html.status_code == 200
    assert 'attachment; filename="glow-workshop-worksheets.html"' in html.headers["Content-Disposition"]
    assert docx.status_code == 200
    assert docx.headers["Content-Type"] == DOCX_MIME
    assert docx.data[:2] == b"PK"


def test_the_guide_and_home_page_point_at_the_pack(client):
    for path in ("/workshop/guide", "/workshop/"):
        page = client.get(path).get_data(as_text=True)
        assert "/workshop/worksheets.docx" in page, path
        assert "/workshop/worksheets.html" in page, path
